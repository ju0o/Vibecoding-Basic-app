from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "release" / "3강-강사용-대본-다듬은본.docx"
ARCHITECTURE_IMAGE = ROOT / "src" / "content" / "assets" / "session-03" / "architecture-evolution-v2.png"
RESTAURANT_IMAGE = ROOT / "src" / "content" / "assets" / "session-03" / "restaurant-cutaway-v2.png"

NAVY = "173149"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "177A72"
GOLD = "9B6C13"
INK = "202833"
MUTED = "66717D"
LIGHT_BLUE = "E8F1F5"
LIGHT_TEAL = "E7F3F0"
LIGHT_GOLD = "FFF5DC"
LINE = "CBD6DE"
WHITE = "FFFFFF"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr()
    fonts = run._element.rPr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run._element.rPr.insert(0, fonts)
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = rgb(color)
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_paragraph_border(paragraph, color=TEAL, size=18):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Slide Title" not in styles:
        slide_style = styles.add_style("Slide Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        slide_style = styles["Slide Title"]
    slide_style.font.name = "Calibri"
    slide_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    slide_style.font.size = Pt(19)
    slide_style.font.bold = True
    slide_style.font.color.rgb = rgb(NAVY)
    slide_style.paragraph_format.space_before = Pt(0)
    slide_style.paragraph_format.space_after = Pt(10)
    slide_style.paragraph_format.keep_with_next = True


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    set_run_font(run, size=9.5, bold=True, color=TEAL)


def add_title(doc, text, size=29, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=NAVY)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    set_run_font(run, size=13.2, color=MUTED)


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [3120, 3120, 3120], indent_dxa=120)
    values = [
        ("12장", "강의 슬라이드"),
        ("30분", "이론 설명"),
        ("30분", "자유 실습"),
    ]
    for cell, (value, label) in zip(table.rows[0].cells, values):
        set_cell_shading(cell, LIGHT_GOLD)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(value)
        set_run_font(r, size=18, bold=True, color=GOLD)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(label)
        set_run_font(r2, size=9.5, color=MUTED)


def add_callout(doc, label, text, fill=LIGHT_TEAL, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.04)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    set_paragraph_border(p, color=color)
    set_paragraph_shading(p, fill)
    label_run = p.add_run(f"{label}  ")
    set_run_font(label_run, size=10, bold=True, color=color)
    text_run = p.add_run(text)
    set_run_font(text_run, size=10.5, color=INK)
    return p


def add_slide_heading(doc, number, title, time):
    doc.add_page_break()
    add_kicker(doc, f"SLIDE {number:02d} · {time}")
    p = doc.add_paragraph(style="Slide Title")
    p.add_run(title)


def add_script(doc, paragraphs):
    doc.add_paragraph("강의 대본", style="Heading 2")
    for text in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = False
        p.add_run(text)


def add_stage(doc, title, text):
    p = doc.add_paragraph(style="Heading 3")
    p.add_run(title)
    body = doc.add_paragraph(text)
    body.paragraph_format.keep_together = False


def add_teacher_note(doc, text):
    add_callout(doc, "강사 메모", text, fill=LIGHT_BLUE, color=BLUE)


def build_document():
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.text = ""
    left = header.add_run("바이브코딩 기초반 · 3강")
    set_run_font(left, size=9, bold=True, color=MUTED)
    header.paragraph_format.space_after = Pt(0)
    add_page_number(section.footer.paragraphs[0])

    add_kicker(doc, "TEACHER SCRIPT · SESSION 03")
    add_title(doc, "개발 용어 이해")
    add_subtitle(doc, "사용자가 작성한 원고를 바탕으로 다듬은 12장 강사용 상세 대본")
    add_metric_strip(doc)

    doc.add_paragraph("이번 강의의 중심 문장", style="Heading 1")
    add_callout(
        doc,
        "핵심",
        "코드를 외우는 시간이 아니라, 화면을 눌렀을 때 서비스의 어느 공간에서 무슨 일이 일어나는지를 이해하는 시간입니다.",
    )
    doc.add_paragraph(
        "수강생이 프론트엔드, 백엔드, 데이터베이스, API라는 단어를 정확히 암기하는 것이 목표는 아닙니다. "
        "각 용어가 맡은 위치를 이해하고, AI에게 수정할 위치와 원하는 결과를 더 구체적으로 말할 수 있으면 충분합니다."
    )
    doc.add_paragraph("진행 원칙", style="Heading 1")
    doc.add_paragraph(
        "애니메이션을 먼저 보여준 뒤 설명합니다. 한 장면이 자동으로 지나가게 두지 않고, 강사가 직접 한 단계씩 눌러 수강생이 관찰할 시간을 줍니다."
    )
    doc.add_paragraph(
        "각 개념은 건축과 식당이라는 익숙한 공간에 먼저 연결합니다. 이후 실제 쇼핑몰 화면과 주문 흐름을 보여주며 개발 용어로 다시 이름을 붙입니다."
    )

    add_slide_heading(doc, 1, "개발의 원리를 눈으로 이해합니다", "약 1분 30초")
    add_callout(doc, "화면 조작", "‘흐름 재생’을 누르고 화면 → 처리 → 저장 → 응답이 차례로 켜지는 장면을 잠시 보여줍니다.")
    add_script(
        doc,
        [
            "안녕하세요. 오늘은 개발 코드를 외우는 시간이 아니라, 우리가 만드는 사이트와 프로그램이 어떤 원리로 움직이는지를 살펴보겠습니다.",
            "우리가 평소 접하는 사이트에는 검색창, 메뉴, 카드, 로그인 버튼, 주문 버튼처럼 눈에 보이는 요소들이 있습니다. 사용자는 화면에서 버튼 하나를 누르지만, 그 뒤에서는 화면이 요청을 보내고, 보이지 않는 곳에서 규칙을 처리하고, 필요한 기록을 저장한 뒤, 결과를 다시 화면으로 돌려주는 일이 이어집니다.",
            "오늘은 이 네 가지 흐름을 중심으로 프론트엔드, 백엔드, 데이터베이스, API가 각각 어디에서 어떤 역할을 하는지 이해해 보겠습니다. 이 원리를 알면 AI가 만든 결과물을 보고도 무엇이 빠졌는지, 어디를 더 요청해야 하는지 판단하기 쉬워집니다.",
        ],
    )
    add_teacher_note(doc, "표지에서는 네 용어를 자세히 설명하지 않습니다. 움직임만 보여주고 ‘이 흐름을 오늘 하나씩 열어보겠습니다’라고 짧게 연결합니다.")

    add_slide_heading(doc, 2, "아이디어가 진짜 서비스가 되는 과정", "약 5분")
    add_callout(doc, "화면 조작", "빈 땅부터 준공까지 일곱 단계를 한 번에 재생하지 말고, 각 버튼을 직접 눌러 설명합니다.")
    doc.add_picture(str(ARCHITECTURE_IMAGE), width=Inches(6.5))
    add_stage(
        doc,
        "1. 빈 땅 · 아이디어",
        "바이브코딩 과정을 하나의 건축 과정으로 비유해 보겠습니다. 처음에는 빈 땅만 있습니다. 만들고 싶은 생각은 있지만 아직 화면도, 파일도, 데이터도 없습니다. 아이디어만 있는 상태입니다.",
    )
    add_stage(
        doc,
        "2. 설계도 · AI와 기획",
        "지난 시간에 GPT, Gemini, Claude 같은 설계 AI와 대화하며 원하는 기능과 디자인을 정리했습니다. 그렇게 정리한 기획과 프롬프트는 건물을 짓기 전에 만드는 설계도와 같습니다. 설계도가 구체적일수록 다음 단계에서 AI IDE가 만들어야 할 범위도 분명해집니다.",
    )
    add_stage(
        doc,
        "3. 철골 · 프로젝트 구조",
        "정리한 프롬프트를 Cursor, Kiro, Codex 같은 AI 코딩 도구에 전달하면 폴더와 파일이 만들어지고 기본 화면이 생깁니다. 아직 사람이 편하게 사용할 수 있는 완성된 서비스는 아니지만, 건물이 무너지지 않도록 뼈대를 세운 것처럼 프로젝트의 구조가 잡힌 상태입니다.",
    )
    add_stage(
        doc,
        "4. 외관 · UI와 화면",
        "AI가 디자인까지 적용하면 메뉴와 버튼, 이미지와 색상이 갖춰지면서 겉으로는 상당히 완성된 서비스처럼 보입니다. 바이브코딩을 처음 시작할 때 가장 만족감이 큰 구간이기도 합니다. 눈에 보이는 결과가 빠르게 생기기 때문입니다.",
    )
    add_stage(
        doc,
        "5. 내부 진입 · 겉만 완성된 상태",
        "하지만 외관이 완성되었다고 실제로 사용할 수 있는 건물은 아닙니다. 안으로 들어가 보니 조명과 수도가 연결되지 않았거나 방의 용도가 정리되지 않았을 수 있습니다. 서비스도 마찬가지입니다. 버튼은 보이지만 눌러도 아무 일도 일어나지 않거나, 로그인과 데이터 저장이 연결되지 않았거나, 오류가 나도 안내가 없을 수 있습니다.",
    )
    add_stage(
        doc,
        "6. 내부 공사 · 세부 기능 조율",
        "이 단계에서 필요한 것이 조율입니다. 조명과 수도, 방 구조와 가구를 하나씩 완성하듯이 버튼의 동작, 로그인, 데이터 저장, 결제, 오류 안내를 하나씩 점검하고 AI에게 다시 요청합니다. AI가 처음부터 모든 것을 알아서 완성해야 하는 것이 아니라, 우리가 필요한 세부 조건을 알려주며 결과물을 조정하는 과정입니다.",
    )
    add_stage(
        doc,
        "7. 준공 · 운영 가능한 서비스",
        "외관과 내부 기능을 모두 확인하고 실제 사용자가 문제없이 이용할 수 있을 때 비로소 준공된 건물처럼 운영 가능한 서비스가 됩니다. 오늘 개발 용어를 배우는 이유도 바로 이 내부 공사를 어디에 요청해야 하는지 알기 위해서입니다.",
    )
    add_callout(doc, "핵심 문장", "우리는 AI가 만든 결과물의 모든 코드를 직접 작성하는 사람이 아니라, 필요한 부분을 발견하고 조율하는 사람입니다.")

    add_slide_heading(doc, 3, "식당 안에서 개발 용어 찾기", "약 3분")
    add_callout(doc, "화면 조작", "주문 카운터 → 주방 → 금고·기록실 → 납품 통로 순서로 누릅니다.")
    doc.add_picture(str(RESTAURANT_IMAGE), width=Inches(6.5))
    add_script(
        doc,
        [
            "이번에는 개발 용어를 식당에 비유해 보겠습니다. 식당 안에는 손님이 보는 공간과 손님에게 보이지 않는 공간이 함께 있습니다.",
            "먼저 손님이 메뉴를 보고 주문하는 카운터가 있습니다. 손님이 직접 보고 선택하고 누르는 곳이기 때문에 프론트엔드와 같습니다.",
            "카운터에서 주문이 들어오면 주방은 정해진 순서와 규칙에 따라 음식을 만듭니다. 주문이 가능한지 확인하고, 재료를 사용하고, 완성된 결과를 내보내는 주방이 백엔드에 해당합니다.",
            "가게의 돈과 회원 기록, 주문 내역을 나중에 다시 찾으려면 금고와 기록실이 필요합니다. 회원, 게시글, 상품, 주문 정보를 보관하는 데이터베이스와 같습니다.",
            "마지막으로 주방이 모든 재료를 직접 만들 수는 없습니다. 외부 거래처에 필요한 재료를 요청하고 납품받는 통로가 필요합니다. 내 서비스가 날씨, 지도, 결제처럼 다른 서비스에 부탁하고 답을 받는 API와 비슷합니다.",
        ],
    )
    add_callout(doc, "정리", "프론트엔드 = 카운터 / 백엔드 = 주방 / 데이터베이스 = 금고와 기록실 / API = 외부 거래처로 이어지는 통로")

    add_slide_heading(doc, 4, "프론트엔드 안에는 무엇이 있을까", "약 3분")
    add_callout(doc, "화면 조작", "검색, 메뉴, 상품 카드, 장바구니, 로그인, 결과 알림을 하나씩 눌러 작은 설명창을 보여줍니다.")
    add_script(
        doc,
        [
            "이제 우리가 가장 자주 접하는 실제 사이트 화면에서 프론트엔드를 조금 더 세밀하게 보겠습니다.",
            "사용자가 원하는 내용을 직접 입력하는 곳은 검색창입니다. 다른 화면으로 이동할 수 있도록 길을 보여주는 것은 메뉴입니다. 상품 사진, 상품명, 가격처럼 같은 규칙으로 반복되는 정보 묶음은 카드라고 부릅니다.",
            "장바구니 담기나 로그인처럼 사용자의 행동을 시작하는 요소가 버튼입니다. 버튼을 눌렀을 때 숫자가 바뀌거나 팝업과 완료 알림이 나타나는 것은 사용자의 행동에 대한 피드백입니다.",
            "프론트엔드는 단순히 예쁜 화면 한 장이 아닙니다. 사용자가 이동하고, 정보를 읽고, 행동하고, 그 결과를 확인하는 모든 접점이 함께 작동하는 공간입니다.",
        ],
    )
    add_teacher_note(doc, "각 요소의 이름을 외우게 하기보다 ‘사용자는 여기에서 무엇을 하는가?’를 먼저 질문합니다.")

    add_slide_heading(doc, 5, "UI와 UX: 보이는 것과 사용되는 방식", "약 3분")
    add_callout(doc, "화면 조작", "‘같은 작업 비교’를 눌러 두 결제 화면에서 커서가 이동하는 횟수와 정보 확인 과정을 비교합니다.")
    add_script(
        doc,
        [
            "프론트엔드를 이야기할 때 UI와 UX라는 단어가 자주 함께 나옵니다. 두 단어는 붙어 다니지만 같은 뜻은 아닙니다.",
            "UI는 버튼의 이름, 색, 크기, 글자, 아이콘, 배치처럼 눈에 보이고 직접 조작하는 요소입니다. UX는 그 요소들을 사용해 목표까지 도달하는 전체 경험입니다. 지금 어디에 있는지, 다음에 무엇을 눌러야 하는지, 최종 결과가 무엇인지 이해하기 쉬운지가 UX에 포함됩니다.",
            "왼쪽 화면도 결제를 할 수는 있지만 현재 단계가 보이지 않고, 확인과 다음처럼 의미가 모호한 버튼이 여러 개 있으며, 배송비와 최종 금액도 늦게 나타납니다. 사용자는 어디로 가야 하는지 여러 번 찾아야 합니다.",
            "오른쪽 화면은 현재 결제 단계, 최종 금액, 다음 행동이 한눈에 보입니다. 같은 목표를 더 적게 찾고, 덜 고민하고, 실수 가능성을 낮추면서 완료할 수 있습니다.",
            "따라서 예쁜 UI가 언제나 좋은 UX를 의미하지는 않습니다. 보기 좋은 화면과 사용하기 쉬운 경험을 함께 설계해야 합니다.",
        ],
    )

    add_slide_heading(doc, 6, "애니메이션은 변화를 설명합니다", "약 3분")
    add_callout(doc, "화면 조작", "여섯 가지 예시를 차례로 누릅니다. 각 예시는 원인 → 변화 → 결과 순서로 설명합니다.")
    add_stage(doc, "메뉴 전환", "메뉴가 갑자기 나타나는 대신 화면 가장자리에서 이어져 들어오면 어디에서 시작된 화면인지 이해하기 쉽습니다.")
    add_stage(doc, "로딩 전환", "처리 중이라는 움직임을 보여주면 사이트가 멈춘 것이 아니라 결과를 기다리고 있다는 사실을 전달할 수 있습니다.")
    add_stage(doc, "장바구니 이동", "상품이 장바구니 쪽으로 이동하면 사용자의 행동과 바뀐 결과가 자연스럽게 연결됩니다.")
    add_stage(doc, "완료 전환", "토스트 알림이나 완료 화면은 작업이 끝났다는 확신과 다음 행동을 알려줍니다.")
    add_stage(doc, "탭 전환", "같은 화면 안에서 현재 보고 있는 영역이 어디인지 움직임으로 표시합니다.")
    add_stage(doc, "입력 검증", "입력이 잘못되었을 때 문제가 있는 위치와 다시 입력해야 할 내용을 바로 보여줍니다.")
    add_callout(doc, "핵심", "좋은 애니메이션은 화려한 장식이 아니라 사용자가 변화의 원인과 결과를 놓치지 않게 돕는 안내입니다.")

    add_slide_heading(doc, 7, "백엔드: 주문 버튼 뒤의 처리실", "약 3분")
    add_callout(doc, "화면 조작", "주문하기를 누른 뒤 ‘다음 처리 보기’를 한 번씩 눌러 접수, 권한, 재고, 결제, 저장, 응답을 설명합니다.")
    add_script(
        doc,
        [
            "백엔드는 사용자의 눈에는 보이지 않지만 서비스의 규칙을 판단하고 실행하는 공간입니다.",
            "쇼핑몰에서 주문하기 버튼을 누르면 먼저 어떤 상품을 몇 개 주문했는지 요청을 접수합니다. 그다음 로그인한 사용자인지, 주문할 권한이 있는지 확인합니다.",
            "재고가 남아 있는지 확인하고, 구매 가능한 상품이라면 결제 승인을 요청합니다. 결제가 성공하면 주문 번호와 구매 내역을 데이터베이스에 저장합니다.",
            "모든 처리가 끝난 뒤에야 주문이 완료되었다는 응답이 브라우저로 돌아옵니다. 사용자에게는 아주 짧은 순간이지만, 화면 뒤에서는 여러 규칙이 순서대로 통과하고 있습니다.",
            "따라서 ‘주문 버튼을 만들어 주세요’라는 요청과 ‘로그인 확인, 재고 확인, 결제, 주문 저장, 완료 응답까지 연결해 주세요’라는 요청은 결과의 깊이가 다릅니다.",
        ],
    )

    add_slide_heading(doc, 8, "데이터베이스: 필요한 기록을 나눠 보관하기", "약 4분")
    add_callout(doc, "화면 조작", "회원가입, 로그인, 게시글 작성, 상품 주문을 수강생에게 먼저 질문한 뒤 각각의 보관함을 누릅니다.")
    add_script(
        doc,
        [
            "데이터베이스는 정보를 한곳에 아무렇게나 쌓아두는 공간이 아니라, 다시 찾기 쉽도록 종류에 맞게 나누어 보관하는 공간입니다.",
            "회원가입을 하면 이름, 이메일, 비밀번호 같은 회원 정보가 회원정보 창고에 새로 저장됩니다. 로그인할 때는 새로 저장하는 것이 아니라, 입력한 정보와 회원정보 창고에 있는 기존 기록을 찾아 비교합니다.",
            "게시글을 작성하면 제목, 내용, 작성자 정보가 게시글 창고에 저장됩니다. 상품을 주문하면 상품 창고에서 재고를 확인하고 수량을 줄인 뒤, 주문자와 상품, 결제 결과를 주문내역 창고에 기록합니다.",
            "이렇게 행동에 따라 저장할 곳과 다시 찾아올 곳이 달라집니다. AI에게 기능을 요청할 때도 어떤 정보를 어디에 저장하고 언제 다시 읽어올지 함께 말하면 결과가 훨씬 명확해집니다.",
        ],
    )
    add_teacher_note(
        doc,
        "개인용 프로그램처럼 정보를 다른 사람과 공유할 필요가 없는 경우에는 브라우저의 로컬 저장소나 컴퓨터의 파일에 기록할 수도 있습니다. 다만 여러 기기와 사용자가 같은 정보를 공유해야 한다면 서버와 데이터베이스가 필요하다고 설명합니다.",
    )

    add_slide_heading(doc, 9, "API: 외부 서비스에 부탁하고 답 받기", "약 3분")
    add_callout(doc, "화면 조작", "‘서울 날씨 불러오기’와 ‘연결 끊기’를 각각 눌러 성공과 실패를 비교합니다.")
    add_script(
        doc,
        [
            "API는 내 서비스와 다른 서비스가 정해진 방식으로 요청과 응답을 주고받는 창구입니다.",
            "예를 들어 내 여행 준비 사이트에서 서울의 현재 날씨를 보여주고 싶다고 해보겠습니다. 내 사이트가 모든 기상 정보를 직접 측정할 수는 없습니다. 날씨 정보를 제공하는 외부 서비스에 ‘서울 날씨를 알려주세요’라고 요청하고, 받은 기온과 하늘 상태를 내 화면에 표시합니다.",
            "결제, 지도, 번역, AI 모델도 같은 방식으로 연결할 수 있습니다. 중요한 점은 상대 서비스가 API를 제공해야 하고, 우리가 그 서비스가 정한 요청 방법을 따라야 한다는 것입니다.",
            "외부 서비스가 API를 제공하지 않거나 연결이 끊기면 원하는 기능을 사용할 수 없습니다. 그래서 API를 사용하는 화면에는 실패했을 때의 안내와 다시 시도하는 방법도 함께 필요합니다.",
        ],
    )

    add_slide_heading(doc, 10, "큰 화면을 역할별로 나누어 요청하기", "약 4분")
    add_callout(doc, "화면 조작", "상단 메뉴, 검색 영역, 상품 카드, 행동 버튼, 결과 알림을 하나씩 누르며 AI 요청 문장이 어떻게 달라지는지 보여줍니다.")
    add_script(
        doc,
        [
            "오늘 배운 용어를 실제 바이브코딩에 어떻게 사용할 수 있는지 정리해 보겠습니다.",
            "예전에는 화면 전체를 보고 ‘이 사이트를 더 좋게 만들어 주세요’라고 한 번에 요청했을 수 있습니다. 이제는 화면을 역할별로 나누어 수정할 위치와 원하는 결과를 더 정확하게 가리킬 수 있습니다.",
            "상단 메뉴에서는 ‘장바구니 숫자를 더 잘 보이게 하고, 모바일에서는 메뉴가 접히게 해 주세요’라고 요청할 수 있습니다.",
            "검색 영역에서는 ‘검색창을 더 크게 만들고, 입력 중인 검색어와 결과가 없을 때의 안내를 보여 주세요’라고 요청할 수 있습니다.",
            "상품 카드에서는 ‘사진, 상품명, 가격, 담기 버튼이 모든 카드에서 같은 위치에 오도록 정리해 주세요’라고 말할 수 있습니다.",
            "행동 버튼에서는 ‘담기 버튼을 누르면 처리 중 표시를 보여주고, 작업이 끝날 때까지 다시 누르지 못하게 해 주세요’라고 요청할 수 있습니다.",
            "결과 알림에서는 ‘상품이 담기면 오른쪽 아래에 완료 알림을 띄우고 장바구니 숫자도 함께 바꿔 주세요’라고 말할 수 있습니다.",
            "용어를 배우는 목적은 어려운 말을 사용하기 위해서가 아닙니다. 내가 원하는 수정 위치와 동작을 AI가 오해하지 않도록 더 구체적으로 전달하기 위해서입니다.",
        ],
    )
    add_callout(doc, "수강생 질문", "여러분 프로젝트에서 지금 바꾸고 싶은 것은 화면, 처리, 저장, 외부 연결 중 어디에 가장 가깝습니까?")

    add_slide_heading(doc, 11, "30분 자유 실습", "30분")
    add_callout(doc, "화면 조작", "타이머 시작을 누릅니다. 별도의 완료 기준을 제시하지 않고 각자 만들고 싶은 프로젝트를 진행하게 합니다.")
    add_script(
        doc,
        [
            "이제 30분 동안 각자 만들고 싶은 프로젝트를 자유롭게 진행하겠습니다.",
            "오늘 배운 용어를 모두 억지로 사용할 필요는 없습니다. 작업하다가 AI에게 설명하기 어려운 부분이 생기면 화면인지, 처리인지, 저장인지, 외부 연결인지 먼저 위치를 나누어 보세요.",
            "중간에 막히거나 AI의 답변이 이해되지 않으면 바로 질문해 주세요. 저는 돌아다니면서 각 프로젝트의 현재 상태와 다음 요청을 함께 정리해 드리겠습니다.",
        ],
    )
    add_teacher_note(doc, "실습 중에는 결과물을 통일하지 않습니다. 수강생마다 프로젝트가 다르므로 요구사항보다 진행 흐름과 AI에게 다시 질문하는 방법을 돕습니다.")

    add_slide_heading(doc, 12, "마무리와 4강 예고", "약 1분")
    add_script(
        doc,
        [
            "오늘은 화면에서 보이는 프론트엔드, 뒤에서 규칙을 처리하는 백엔드, 기록을 보관하는 데이터베이스, 외부 서비스와 연결하는 API를 살펴봤습니다.",
            "모든 용어를 정확하게 외울 필요는 없습니다. 내가 수정하려는 기능이 어느 공간에 가까운지 구분하고 AI에게 설명할 수 있으면 충분합니다.",
            "다음 시간에는 오늘 화면에서 보았던 버튼과 기능이 실제 프로젝트의 어느 폴더와 파일에 들어 있는지 찾아보겠습니다. 화면의 결과와 파일 구조를 연결하면 AI가 어떤 파일을 만들고 수정했는지도 조금씩 읽을 수 있게 됩니다.",
        ],
    )
    add_callout(doc, "다음 강의", "4강 · 파일 구조 이해: 폴더의 역할, 수정할 파일 찾기, AI가 바꾼 위치 확인하기")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
